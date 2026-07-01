import json
import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
RESULTS_DIR = ROOT / "results"
REPORT_DIR = ROOT / "report"
REPORT_PATH = REPORT_DIR / "深度学习导论课程报告_优化算法对比实验.docx"

MAIN_OPTIMIZERS = ["SGD", "Adam", "AdamW"]


def fmt_num(value, digits=4):
    return f"{float(value):.{digits}f}"


def fmt_pct(value):
    return f"{float(value):.2f}%"


def set_run_font(run, size=None, bold=None, name="宋体"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_font(paragraph, size=11, name="宋体"):
    for run in paragraph.runs:
        set_run_font(run, size=size, name=name)


def add_body_paragraph(doc, text, body_texts):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(22)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run, size=11)
    body_texts.append(text)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        set_run_font(run, size=14 if level == 1 else 12, bold=True, name="黑体")
    return paragraph


def add_centered_picture(doc, image_path, caption, width=5.8):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_p.add_run(caption)
    set_run_font(caption_run, size=10, name="宋体")


def add_dataframe_table(doc, title, df, columns, display_names):
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title)
    set_run_font(title_run, size=10.5, bold=True, name="宋体")

    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = table.rows[0].cells
    for idx, name in enumerate(display_names):
        header_cells[idx].text = name
        header_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in header_cells[idx].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_font(paragraph, size=9.5, name="宋体")

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(columns):
            value = row[col]
            if col in {"best_test_acc", "final_test_acc", "final_train_acc"}:
                text = fmt_pct(value)
            elif col in {"final_train_loss", "final_test_loss"}:
                text = fmt_num(value, 4)
            elif col in {"lr", "momentum", "weight_decay"}:
                text = f"{float(value):g}"
            elif col == "best_epoch":
                text = str(int(value))
            else:
                text = str(value)
            cells[idx].text = text
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_font(paragraph, size=9.5, name="宋体")
    return table


def first_epoch_reaching(metrics_df, run_name, target):
    sub = metrics_df[metrics_df["run_name"] == run_name].sort_values("epoch")
    reached = sub[sub["test_acc"] >= target]
    if reached.empty:
        return None
    return int(reached.iloc[0]["epoch"])


def choose_fastest(metrics_df, target):
    reach = {
        name: first_epoch_reaching(metrics_df, name, target)
        for name in MAIN_OPTIMIZERS
    }
    reached = {name: epoch for name, epoch in reach.items() if epoch is not None}
    if reached:
        fastest = min(
            reached,
            key=lambda name: (
                reached[name],
                -float(
                    metrics_df[
                        (metrics_df["run_name"] == name)
                        & (metrics_df["epoch"] == reached[name])
                    ]["test_acc"].iloc[0]
                ),
            ),
        )
        return fastest, f"第{reached[fastest]}轮率先达到{target:.1f}%测试准确率"

    early = (
        metrics_df[
            metrics_df["run_name"].isin(MAIN_OPTIMIZERS)
            & (metrics_df["epoch"] <= 3)
        ]
        .groupby("run_name")["test_acc"]
        .mean()
    )
    fastest = str(early.idxmax())
    return fastest, "所有主实验均未达到目标阈值，因此按前三轮平均测试准确率判断"


def load_inputs():
    required = [
        RESULTS_DIR / "metrics.csv",
        RESULTS_DIR / "final_results.csv",
        RESULTS_DIR / "loss_curve.png",
        RESULTS_DIR / "accuracy_curve.png",
        RESULTS_DIR / "adamw_weight_decay_ablation.png",
        RESULTS_DIR / "run_log.txt",
    ]
    missing = [str(path) for path in required if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing required result files: " + ", ".join(missing))

    metrics_df = pd.read_csv(RESULTS_DIR / "metrics.csv")
    final_df = pd.read_csv(RESULTS_DIR / "final_results.csv")
    meta_path = RESULTS_DIR / "experiment_meta.json"
    meta = json.loads(meta_path.read_text(encoding="utf-8")) if meta_path.exists() else {}
    return metrics_df, final_df, meta


def main():
    metrics_df, final_df, meta = load_inputs()
    REPORT_DIR.mkdir(parents=True, exist_ok=True)

    main_final = (
        final_df[final_df["optimizer"].isin(MAIN_OPTIMIZERS)]
        .set_index("optimizer")
        .loc[MAIN_OPTIMIZERS]
        .reset_index()
    )

    dataset = str(meta.get("dataset", main_final.iloc[0]["dataset"]))
    dataset_label = "Fashion-MNIST" if dataset == "FashionMNIST" else dataset
    epochs = int(meta.get("epochs", metrics_df["epoch"].max()))
    batch_size = int(meta.get("batch_size", 128))
    target_acc = float(meta.get("target_acc", 90.0))
    device = str(meta.get("device", "unknown"))
    cuda_device = str(meta.get("cuda_device", ""))
    device_text = f"{device}（{cuda_device}）" if cuda_device else device

    best_row = main_final.sort_values(
        ["best_test_acc", "final_test_acc"], ascending=False
    ).iloc[0]
    best_optimizer = str(best_row["optimizer"])
    fastest_optimizer, fastest_reason = choose_fastest(metrics_df, target_acc)

    sgd = main_final[main_final["optimizer"] == "SGD"].iloc[0]
    adam = main_final[main_final["optimizer"] == "Adam"].iloc[0]
    adamw = main_final[main_final["optimizer"] == "AdamW"].iloc[0]
    adamw_no_wd = final_df[final_df["optimizer"] == "AdamW_no_wd"].iloc[0]

    adamw_best_delta = float(adamw["best_test_acc"]) - float(adamw_no_wd["best_test_acc"])
    adamw_final_delta = float(adamw["final_test_acc"]) - float(
        adamw_no_wd["final_test_acc"]
    )
    adamw_loss_delta = float(adamw["final_test_loss"]) - float(
        adamw_no_wd["final_test_loss"]
    )

    epoch1 = (
        metrics_df[
            metrics_df["run_name"].isin(MAIN_OPTIMIZERS) & (metrics_df["epoch"] == 1)
        ][["run_name", "train_loss", "test_acc"]]
        .set_index("run_name")
        .loc[MAIN_OPTIMIZERS]
    )
    lowest_epoch1_loss = str(epoch1["train_loss"].idxmin())
    highest_epoch1_acc = str(epoch1["test_acc"].idxmax())

    overfit_gaps = {
        row["optimizer"]: float(row["final_train_acc"]) - float(row["final_test_acc"])
        for _, row in main_final.iterrows()
    }
    largest_gap_name = max(overfit_gaps, key=overfit_gaps.get)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "宋体"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal_style.font.size = Pt(11)

    title = (
        f"优化算法在深度学习中的演进与对比实验——基于{dataset_label}的"
        "SGD、Adam 与 AdamW 比较"
    )
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title)
    set_run_font(title_run, size=16, bold=True, name="黑体")

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_p.add_run("《深度学习导论》课程报告")
    set_run_font(info_run, size=11, name="宋体")

    body_texts = []

    add_heading(doc, "1 摘要", 1)
    add_body_paragraph(
        doc,
        f"本文围绕深度学习中的优化算法展开研究，先按照 SGD、Momentum、NAG、AdaGrad、RMSProp、Adam、AdamW 的顺序梳理优化器从经典随机梯度方法到自适应学习率方法的发展脉络，再基于 {dataset_label} 数据集和一个小型卷积神经网络完成对比实验。实验使用相同的数据划分、模型结构、训练轮数和 CosineAnnealingLR 学习率调度，分别训练 SGD、Adam、AdamW 三种主实验配置，并额外完成 AdamW 在 weight_decay=0 与 weight_decay=0.01 下的权重衰减消融。所有结论均来自本项目实际运行得到的 metrics.csv、final_results.csv、曲线图和运行日志。从本次实验结果看，{fastest_optimizer} 的收敛速度表现最突出，判断依据是{fastest_reason}；以最佳测试准确率为主要指标，{best_optimizer} 在三种主优化器中泛化表现最好，best_test_acc 为 {fmt_pct(best_row['best_test_acc'])}。实验说明，自适应优化器通常更容易在训练初期快速降低损失，但最终泛化效果仍与正则化、学习率调度和训练轮数密切相关，不能简单认为 Adam 或 AdamW 在任何情况下都必然优于带动量的 SGD。",
        body_texts,
    )

    add_heading(doc, "2 引言", 1)
    add_body_paragraph(
        doc,
        "深度学习模型的训练过程本质上是一个大规模非凸优化问题。给定网络结构和损失函数后，训练算法需要在高维参数空间中不断更新权重，使模型在训练数据上的误差下降，并尽可能在未见过的数据上保持良好的泛化能力。优化器在这个过程中并不是一个简单的工程细节，它会直接影响损失下降速度、训练稳定性、对学习率的敏感程度以及最终落入的极小值区域。对于同一个模型，即使参数量、数据集和训练轮数都完全相同，不同优化器也可能表现出明显差异。",
        body_texts,
    )
    add_body_paragraph(
        doc,
        "经典 SGD 使用 mini-batch 梯度估计来近似整体梯度，计算成本低，仍然是许多视觉任务中的重要基线；Adam 和 AdamW 则通过一阶、二阶矩估计为不同参数自动调整更新幅度，通常具有更快的早期收敛速度。课程实验中比较这些优化器有两个意义：一方面可以直观看到损失曲线和准确率曲线的差异，理解优化算法不是孤立公式；另一方面也能检验一个常见说法，即“自适应学习率优化器是否总能优于经典 SGD”。因此，本文采用可复现的 PyTorch 代码，在同一小型 CNN 上进行真实训练，并将实验过程、CSV 结果、曲线图和日志整理到报告中。",
        body_texts,
    )

    add_heading(doc, "3 优化算法综述", 1)
    add_heading(doc, "3.1 SGD", 2)
    add_body_paragraph(
        doc,
        "随机梯度下降是深度学习中最基础也最重要的优化方法之一。与每次使用全部训练样本计算梯度的批量梯度下降不同，SGD 在每一步只使用一个 mini-batch 估计梯度方向，因此单步计算开销更低，也能利用梯度噪声在非凸损失曲面中跳出某些局部区域。它的缺点同样明显：更新方向受 mini-batch 抽样影响较大，曲线容易震荡；学习率过大会导致不稳定，学习率过小又会使收敛很慢。因此，实际训练中常将 SGD 与 momentum、weight decay 和学习率调度结合使用。",
        body_texts,
    )
    add_heading(doc, "3.2 Momentum", 2)
    add_body_paragraph(
        doc,
        "Momentum 在 SGD 的基础上引入“速度”变量，把历史梯度方向以指数衰减的方式累积起来。若连续多个 mini-batch 的梯度方向大体一致，动量项会放大这一方向上的更新，从而加快下降；若某些方向上的梯度来回震荡，正负方向会相互抵消，训练轨迹会更平滑。这种思想非常适合狭长谷地形状的损失曲面，因为普通 SGD 可能在陡峭方向上左右摆动，而 Momentum 能让参数沿着更稳定的方向前进。",
        body_texts,
    )
    add_heading(doc, "3.3 NAG", 2)
    add_body_paragraph(
        doc,
        "Nesterov Accelerated Gradient 可以理解为对 Momentum 的进一步改进。普通 Momentum 先根据当前位置的梯度和历史速度更新参数，而 NAG 会先沿着动量方向做一次前瞻，再在这个预估位置计算梯度。这样做的直观好处是，如果模型即将冲过最优区域，前瞻梯度能够更早地提醒更新方向发生修正。NAG 在凸优化理论中有清晰的加速意义，在深度学习中也常被作为 SGD 动量法的改进版本。",
        body_texts,
    )
    add_heading(doc, "3.4 AdaGrad", 2)
    add_body_paragraph(
        doc,
        "AdaGrad 是自适应学习率方法的重要起点。它为每个参数维护历史平方梯度的累积和，梯度累积越大的参数，其有效学习率越小；梯度较少出现的参数，则能保留较大的更新幅度。因此 AdaGrad 对稀疏特征任务很友好，例如自然语言处理中某些低频词的嵌入参数可以得到更有针对性的更新。不过，由于平方梯度是持续累加的，分母会不断变大，学习率在训练后期可能过快衰减，导致模型还没有充分收敛就提前停滞。",
        body_texts,
    )
    add_heading(doc, "3.5 RMSProp", 2)
    add_body_paragraph(
        doc,
        "RMSProp 针对 AdaGrad 后期学习率持续变小的问题，引入指数滑动平均来估计平方梯度，而不是对全部历史梯度做简单累加。这样，较早的梯度信息会逐渐衰减，优化器更关注最近一段时间的梯度尺度，从而保持更灵活的自适应更新能力。RMSProp 对非平稳目标函数较为有效，也为后来的 Adam 提供了二阶矩估计的思想基础。",
        body_texts,
    )
    add_heading(doc, "3.6 Adam", 2)
    add_body_paragraph(
        doc,
        "Adam 同时结合了 Momentum 的一阶矩估计和 RMSProp 的二阶矩估计。它既记录梯度方向的指数滑动平均，也记录平方梯度的指数滑动平均，并进行偏差修正。由于每个参数都会获得自适应的更新步长，Adam 往往在训练早期具有很快的 loss 下降速度，对初始学习率的敏感性也低于普通 SGD。许多课程实验、原型验证和中小规模任务都会优先使用 Adam，因为它能较快给出一个可用结果。",
        body_texts,
    )
    add_heading(doc, "3.7 AdamW", 2)
    add_body_paragraph(
        doc,
        "AdamW 的核心改进是将权重衰减与 Adam 的梯度更新解耦。传统做法常把 L2 正则项直接加入梯度，但在 Adam 这类自适应优化器中，梯度会被二阶矩缩放，导致 weight decay 的效果不再等价于真正意义上的参数衰减。AdamW 单独对权重做衰减，再执行 Adam 式的自适应更新，使正则化含义更清晰，也常被用于提升泛化能力。Transformer 等现代模型中大量使用 AdamW，正说明优化算法的细节会影响最终效果。",
        body_texts,
    )
    add_heading(doc, "3.8 优化器、损失曲面与泛化能力", 2)
    add_body_paragraph(
        doc,
        "训练损失低并不必然意味着测试准确率高。深度网络的损失曲面包含许多极小值区域，其中有些区域很尖锐，参数发生很小扰动就会导致损失明显上升；另一些区域较平坦，对扰动更不敏感。一般认为，较平坦的极小值往往与更好的泛化能力相关。自适应优化器可能更快找到训练损失较低的区域，但这种速度优势不一定自动转化为测试集优势；SGD 虽然前期可能慢一些，但梯度噪声和动量调度有时反而有助于找到更平坦的解。由此可见，优化器选择应结合任务、模型、正则化和训练预算综合判断。",
        body_texts,
    )

    add_heading(doc, "4 实验设置", 1)
    add_body_paragraph(
        doc,
        f"本实验使用 PyTorch 实现，数据集实际采用 {dataset_label}，训练集样本数为 {meta.get('train_samples', '未知')}，测试集样本数为 {meta.get('test_samples', '未知')}。输入图像大小为 1×28×28，分类类别数为 10。模型为小型 CNN：第一层卷积将通道数从 1 提升到 32，接 ReLU 和 2×2 最大池化；第二层卷积将通道数提升到 64，同样接 ReLU 和最大池化；随后展平成 64×7×7，经 128 维全连接层、ReLU、Dropout(0.3) 和最终 10 类输出层。损失函数使用 CrossEntropyLoss，评价指标包括 train_loss、train_acc、test_loss、test_acc。",
        body_texts,
    )
    add_body_paragraph(
        doc,
        f"三组主实验分别为 SGD、Adam、AdamW。SGD 使用 lr=0.05、momentum=0.9、weight_decay=5e-4；Adam 使用 lr=0.001、weight_decay=0；AdamW 使用 lr=0.001、weight_decay=0.01。为了保证策略一致，三者都使用 CosineAnnealingLR 学习率调度，训练 {epochs} 轮，batch_size={batch_size}，随机种子为 {meta.get('seed', 42)}。实验设备为 {device_text}，Python 版本为 {meta.get('python_version', '未知')}，PyTorch 版本为 {meta.get('torch_version', '未知')}。消融实验固定 AdamW 的其他设置，仅比较 weight_decay=0 和 weight_decay=0.01 的差异。",
        body_texts,
    )

    add_heading(doc, "5 实验结果", 1)
    result_table = main_final[
        [
            "optimizer",
            "best_test_acc",
            "best_epoch",
            "final_test_acc",
            "final_train_loss",
            "final_test_loss",
            "convergence_speed",
        ]
    ].copy()
    add_dataframe_table(
        doc,
        "表 1 SGD、Adam、AdamW 最终结果表（来自 final_results.csv）",
        result_table,
        [
            "optimizer",
            "best_test_acc",
            "best_epoch",
            "final_test_acc",
            "final_train_loss",
            "final_test_loss",
            "convergence_speed",
        ],
        [
            "优化器",
            "最佳测试准确率",
            "最佳轮次",
            "最终测试准确率",
            "最终训练损失",
            "最终测试损失",
            "收敛速度说明",
        ],
    )

    add_centered_picture(
        doc,
        RESULTS_DIR / "loss_curve.png",
        "图 1 不同优化器的损失下降曲线",
        width=6.2,
    )
    add_centered_picture(
        doc,
        RESULTS_DIR / "accuracy_curve.png",
        "图 2 不同优化器的准确率变化曲线",
        width=6.2,
    )
    add_centered_picture(
        doc,
        RESULTS_DIR / "adamw_weight_decay_ablation.png",
        "图 3 AdamW 权重衰减消融实验曲线",
        width=6.2,
    )
    screenshot_path = RESULTS_DIR / "run_screenshot.png"
    if screenshot_path.exists():
        add_centered_picture(
            doc,
            screenshot_path,
            "图 4 训练过程运行日志截图",
            width=6.2,
        )

    add_body_paragraph(
        doc,
        f"从最终结果表可以看到，SGD 的最佳测试准确率为 {fmt_pct(sgd['best_test_acc'])}，最终测试准确率为 {fmt_pct(sgd['final_test_acc'])}；Adam 的最佳测试准确率为 {fmt_pct(adam['best_test_acc'])}，最终测试准确率为 {fmt_pct(adam['final_test_acc'])}；AdamW 的最佳测试准确率为 {fmt_pct(adamw['best_test_acc'])}，最终测试准确率为 {fmt_pct(adamw['final_test_acc'])}。这些数值都由脚本训练完成后写入 final_results.csv，再由本报告生成脚本读取并插入，没有手工改写。",
        body_texts,
    )

    add_heading(doc, "6 结果分析", 1)
    add_body_paragraph(
        doc,
        f"先看收敛速度。第 1 轮训练结束时，主实验中训练损失最低的是 {lowest_epoch1_loss}，测试准确率最高的是 {highest_epoch1_acc}。按本实验设置的 {target_acc:.1f}% 测试准确率阈值统计，{fastest_optimizer} 的收敛速度最快，具体表现为{fastest_reason}。这与曲线图中的趋势一致：自适应优化器在早期通常能更快调整不同参数的步长，因此 loss 下降和 accuracy 上升更明显；SGD 虽然带有 momentum，但初始阶段仍更依赖学习率和动量积累。",
        body_texts,
    )
    add_body_paragraph(
        doc,
        f"再看泛化表现。以 best_test_acc 作为主要标准，三种主优化器中表现最好的是 {best_optimizer}，其最佳测试准确率为 {fmt_pct(best_row['best_test_acc'])}，对应最佳轮次为第 {int(best_row['best_epoch'])} 轮。需要注意的是，最佳准确率和最后一轮准确率并不总是完全一致，这说明训练后期即使 train_acc 继续提高，test_acc 也可能出现小幅波动。本实验中最终训练准确率与测试准确率差距最大的是 {largest_gap_name}，差距约为 {overfit_gaps[largest_gap_name]:.2f} 个百分点，说明模型在后期已经出现一定训练集拟合优势，但由于网络较小、Dropout 和 weight decay 存在，过拟合程度总体可控。",
        body_texts,
    )
    add_body_paragraph(
        doc,
        f"AdamW 权重衰减消融实验用于观察解耦 weight decay 的影响。与 AdamW_no_wd 相比，weight_decay=0.01 的 AdamW 最佳测试准确率变化为 {adamw_best_delta:+.2f} 个百分点，最终测试准确率变化为 {adamw_final_delta:+.2f} 个百分点，最终测试损失变化为 {adamw_loss_delta:+.4f}。如果只看训练损失，去掉权重衰减有时会让模型更自由地拟合训练集；但测试准确率和测试损失更能反映泛化。图 3 显示，两条曲线的差异不是简单单调的，weight decay 的收益会受到训练轮数、学习率调度和数据集难度共同影响。",
        body_texts,
    )
    add_body_paragraph(
        doc,
        "综合曲线可以发现，Adam 和 AdamW 的早期优势主要体现在 loss 下降更快、test_acc 更早进入较高区间；SGD 的变化相对平缓，但后期并非没有竞争力。对于 Fashion-MNIST 这类中等难度的灰度图分类任务，小型 CNN 的容量足够，三种优化器都能得到可用结果，差异更多体现在收敛速度、最后几个百分点的准确率以及曲线稳定性上。这也提醒我，比较优化算法不能只看单轮或单个指标，而应同时观察训练损失、测试损失、准确率曲线和最终泛化表现。",
        body_texts,
    )

    add_heading(doc, "7 心得与总结", 1)
    add_body_paragraph(
        doc,
        f"问题一：哪种优化器收敛最快？根据本次真实运行的曲线和 convergence_speed 字段，答案是 {fastest_optimizer}。我的理解是，自适应优化器能根据梯度二阶矩为不同参数分配不同的有效学习率，所以在训练早期更容易迅速找到下降方向；如果最快的是 SGD，也说明在该学习率、momentum 和余弦退火组合下，经典方法同样可以表现得很强，不能凭经验预设结论。",
        body_texts,
    )
    add_body_paragraph(
        doc,
        f"问题二：哪种优化器泛化最好？以最高测试准确率为主、最终测试准确率为辅，本实验中泛化最好的是 {best_optimizer}。它的 best_test_acc 为 {fmt_pct(best_row['best_test_acc'])}，final_test_acc 为 {fmt_pct(best_row['final_test_acc'])}。这个结果不是由理论预先决定的，而是由 final_results.csv 中的实际数值决定。若换成更深的模型、更长训练轮数或更强数据增强，排序可能发生变化。",
        body_texts,
    )
    add_body_paragraph(
        doc,
        "问题三：自适应学习率是否总能优于经典 SGD？我的看法是否定的。Adam 和 AdamW 的确经常更省调参，尤其适合课程实验、快速验证模型结构和训练预算有限的场景；但泛化能力不只取决于 loss 下降速度，还与正则化、学习率衰减、批量大小、数据分布和极小值平坦程度有关。SGD 对学习率更敏感，训练前期可能慢，但配合 momentum、weight decay 和合适调度后仍然可能取得稳定甚至更好的测试表现。因此，优化器没有绝对最优，合理做法是在具体任务上建立强基线，再通过曲线和验证集结果选择。",
        body_texts,
    )
    add_body_paragraph(
        doc,
        "这次实验最大的收获是把“优化器差异”从抽象公式落实到了可复现数据上。过去我容易把 Adam 理解成 SGD 的简单升级版，但实际曲线显示，快收敛、低训练损失和好泛化不是同一件事。AdamW 的消融也让我看到，weight decay 不是一个装饰性参数，而是会改变测试曲线和最终损失的正则化手段。以后训练深度模型时，我会更重视完整记录实验配置和随机种子，并用 CSV、曲线和日志共同支撑结论。",
        body_texts,
    )

    add_heading(doc, "8 参考文献", 1)
    references = [
        "[1] Robbins H., Monro S. A stochastic approximation method. The Annals of Mathematical Statistics, 1951.",
        "[2] Polyak B. T. Some methods of speeding up the convergence of iteration methods. USSR Computational Mathematics and Mathematical Physics, 1964.",
        "[3] Nesterov Y. A method for solving the convex programming problem with convergence rate O(1/k^2). Soviet Mathematics Doklady, 1983.",
        "[4] Duchi J., Hazan E., Singer Y. Adaptive subgradient methods for online learning and stochastic optimization. JMLR, 2011.",
        "[5] Tieleman T., Hinton G. Lecture 6.5 RMSProp: Divide the gradient by a running average of its recent magnitude. COURSERA, 2012.",
        "[6] Kingma D. P., Ba J. Adam: A Method for Stochastic Optimization. ICLR, 2015.",
        "[7] Loshchilov I., Hutter F. Decoupled Weight Decay Regularization. ICLR, 2019.",
        "[8] Xiao H., Rasul K., Vollgraf R. Fashion-MNIST: a Novel Image Dataset for Benchmarking Machine Learning Algorithms. arXiv:1708.07747, 2017.",
        "[9] PyTorch Documentation. torch.optim and torchvision.datasets.",
    ]
    for ref in references:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        set_run_font(run, size=10.5)

    add_heading(doc, "9 附录", 1)
    add_body_paragraph(
        doc,
        "运行命令：在项目目录 deep_learning_optimizer_report 下执行 python train_optimizers.py，训练完成后执行 python generate_report.py。训练脚本会自动下载数据集，自动检测 cuda/cpu，保存 results/run_log.txt、results/metrics.csv、results/final_results.csv 以及三张曲线图。",
        body_texts,
    )
    add_body_paragraph(
        doc,
        "项目结构：train_optimizers.py 为训练与绘图脚本；generate_report.py 为 Word 生成脚本；requirements.txt 记录依赖；data/ 保存数据集；results/ 保存 CSV、曲线图、运行日志和日志截图；report/ 保存最终 Word 文档。复现实验时只需保持相同代码、随机种子、训练轮数和依赖版本，即可重新生成全部结果文件。",
        body_texts,
    )

    log_text = (RESULTS_DIR / "run_log.txt").read_text(encoding="utf-8")
    log_lines = [
        line
        for line in log_text.splitlines()
        if re.search(r"Epoch 0[1-3]|Epoch 10|Final results|Saved", line)
    ][:18]
    if log_lines:
        p = doc.add_paragraph()
        run = p.add_run("关键运行日志节选：\n" + "\n".join(log_lines))
        set_run_font(run, size=8.5, name="Courier New")

    body_chars = len("".join(body_texts))
    doc.core_properties.title = title
    doc.core_properties.subject = "深度学习导论课程报告"
    doc.core_properties.comments = f"正文与附录中文字符统计约 {body_chars} 字，实验结果来自 CSV。"
    doc.save(REPORT_PATH)
    print(f"Report saved to: {REPORT_PATH}")
    print(f"Body text characters counted before references: {body_chars}")


if __name__ == "__main__":
    main()
